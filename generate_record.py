import os
import subprocess
import requests
from docx import Document
from docx.shared import Pt

# === CONFIG ===
BASE_DIR = "C++/Lab"
LABS_TO_PROCESS = ["Lab - 1"]
DOC_FILE = "Lab_Record.docx"
AI_ENDPOINT = "https://api-inference.huggingface.co/models/google/flan-t5-small"
HEADERS = {"Authorization": "Bearer "}  # Leave empty for public inference

# === FUNCTIONS ===

def generate_algorithm(code):
    prompt = f"Generate an algorithm in lab report format for the following C++ code. Include Input, Output, and numbered steps:\n\n{code}"
    response = requests.post(AI_ENDPOINT, headers=HEADERS, json={"inputs": prompt})
    try:
        return response.json()[0]['generated_text']
    except:
        return "[Algorithm generation failed]"

def compile_and_run_cpp(file_path):
    executable = file_path.replace(".cpp", "")
    try:
        subprocess.run(["g++", file_path, "-o", executable], check=True, capture_output=True)
        result = subprocess.run([f"./{executable}"], capture_output=True, text=True, input="", timeout=5)
        return result.stdout.strip()
    except subprocess.CalledProcessError as e:
        return f"[Compilation error]\n{e.stderr}"
    except Exception as e:
        return f"[Runtime error] {e}"

def add_section(doc, heading, text, bold=False):
    doc.add_heading(heading, level=2)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True

# === MAIN SCRIPT ===

document = Document()
document.add_heading('C++ Lab Record', 0)

for lab in LABS_TO_PROCESS:
    lab_path = os.path.join(BASE_DIR, lab)
    document.add_heading(lab, level=1)
    
    for file in os.listdir(lab_path):
        if file.endswith(".cpp"):
            file_path = os.path.join(lab_path, file)
            with open(file_path, "r") as f:
                code = f.read()

            document.add_paragraph(f"\nFile: {file}", style='Intense Quote')

            # 1. Algorithm
            algorithm = generate_algorithm(code)
            add_section(document, "Algorithm:", algorithm)

            # 2. Code
            add_section(document, "Code:", code)

            # 3. Output
            output = compile_and_run_cpp(file_path)
            add_section(document, "Output:", output)

document.save(DOC_FILE)
print(f"✅ Lab record saved to {DOC_FILE}")
