from docx import Document
from docx.shared import Pt

def create_lab_record(filename="C++_Lab_Record.docx", total_labs=12):
    doc = Document()

    for i in range(1, total_labs + 1):
        # Create custom heading for "Lab - number"
        heading = doc.add_paragraph()
        run = heading.add_run(f'Lab - {i}')
        run.bold = True
        run.font.size = Pt(20)  # You can increase this for larger size

        doc.add_paragraph()  # Add space after heading

        sections = [
            "Aim",
            "Algorithms",
            "Input",
            "Output",
            "Code",
            "Output"
        ]

        for section in sections:
            doc.add_heading(f'{section}:', level=2)
            doc.add_paragraph()

        if i != total_labs:
            doc.add_page_break()

    doc.save(filename)
    print(f"Document '{filename}' created successfully.")

create_lab_record()
